import argparse
import json
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as RLImage, KeepTogether, Paragraph, SimpleDocTemplate
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    from PIL import Image as PILImage
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False


LAYOUT_PATH = Path(__file__).resolve().parent.parent / "shared" / "disputeLetterLayout.json"
LAYOUT = json.loads(LAYOUT_PATH.read_text())

CONTENT_WIDTH_INCHES = LAYOUT["pageWidthInches"] - 2 * LAYOUT["marginInches"]
MAX_IMAGE_HEIGHT_INCHES = 8.0  # leave room for the caption inside the content frame


def scaled_image_size(image_path: Path) -> dict:
    """300-DPI exhibit PNG -> printed inches, capped to the content frame.
    Shared sizing for letter figures and the memorandum (which imports it
    from here — dependency direction is memorandum -> letter generator)."""
    with PILImage.open(image_path) as image:
        width_px, height_px = image.size
    width_in = min(CONTENT_WIDTH_INCHES, width_px / 300.0)
    height_in = height_px * (width_in / width_px)
    if height_in > MAX_IMAGE_HEIGHT_INCHES:
        scale = MAX_IMAGE_HEIGHT_INCHES / height_in
        width_in *= scale
        height_in = MAX_IMAGE_HEIGHT_INCHES
    return {"width": width_in, "height": height_in}


@dataclass
class InlineRun:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    line_break: bool = False


@dataclass
class Block:
    kind: str
    runs: List[InlineRun]
    css_class: str = ""
    image_src: str = ""


class RichTextBlockParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.blocks: List[Block] = []
        self.current_runs: List[InlineRun] = []
        self.current_kind: Optional[str] = None
        self.current_css_class: str = ""
        self.list_stack: List[str] = []
        self.bold = False
        self.italic = False
        self.underline = False

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        attrs_map = {name.lower(): value for name, value in attrs}
        if tag in {"p", "h2", "h3"}:
            self._flush_block()
            self.current_kind = "p"
            self.current_css_class = attrs_map.get("class", "")
            self.current_runs = []
        elif tag in {"ul", "ol"}:
            self.list_stack.append(tag)
        elif tag == "li":
            self._flush_block()
            self.current_kind = "p"
            self.current_css_class = attrs_map.get("class", "")
            self.current_runs = []
        elif tag in {"strong", "b"}:
            self.bold = True
        elif tag in {"em", "i"}:
            self.italic = True
        elif tag == "u":
            self.underline = True
        elif tag == "br":
            self.current_runs.append(InlineRun(text="", line_break=True))
        elif tag == "img":
            # Exhibit figures (letter Mode A). Appended directly: the text-only
            # gate in _flush_block would drop a runless block. handle_startendtag
            # delegates here, so <img> and <img/> both land in this branch.
            src = attrs_map.get("data-exhibit-file") or attrs_map.get("src") or ""
            self._flush_block()
            if src:
                self.blocks.append(
                    Block(kind="image", runs=[], css_class=attrs_map.get("class", ""), image_src=src)
                )

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in {"p", "h2", "h3", "li"}:
            self._flush_block()
        elif tag in {"ul", "ol"}:
            if self.list_stack:
                self.list_stack.pop()
        elif tag in {"strong", "b"}:
            self.bold = False
        elif tag in {"em", "i"}:
            self.italic = False
        elif tag == "u":
            self.underline = False

    def handle_data(self, data):
        text = data.replace("\xa0", " ")
        if not text:
            return
        self.current_runs.append(
            InlineRun(
                text=text,
                bold=self.bold,
                italic=self.italic,
                underline=self.underline,
            )
        )

    def _flush_block(self):
        if self.current_kind and any(run.text.strip() or run.line_break for run in self.current_runs):
            self.blocks.append(Block(kind=self.current_kind, runs=list(self.current_runs), css_class=self.current_css_class))
        self.current_kind = None
        self.current_css_class = ""
        self.current_runs = []

    def finalize(self):
        self._flush_block()
        return self.blocks


def parse_blocks(html: str) -> List[Block]:
    parser = RichTextBlockParser()
    parser.feed(html)
    return parser.finalize()


def configure_document(document: Document):
    section = document.sections[0]
    section.page_height = Inches(LAYOUT["pageHeightInches"])
    section.page_width = Inches(LAYOUT["pageWidthInches"])
    section.top_margin = Inches(LAYOUT["marginInches"])
    section.bottom_margin = Inches(LAYOUT["marginInches"])
    section.left_margin = Inches(LAYOUT["marginInches"])
    section.right_margin = Inches(LAYOUT["marginInches"])
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(LAYOUT["fontSizePt"])


def add_runs_to_paragraph(paragraph, runs: List[InlineRun]):
    for run_data in runs:
        if run_data.line_break:
            paragraph.add_run().add_break()
            continue
        run = paragraph.add_run(run_data.text)
        run.bold = run_data.bold
        run.italic = run_data.italic
        run.underline = run_data.underline
        run.font.name = "Times New Roman"
        run.font.size = Pt(LAYOUT["fontSizePt"])


def apply_block_spacing(paragraph, block: Block):
    css_classes = set(block.css_class.split())
    paragraph.paragraph_format.line_spacing = LAYOUT["lineHeight"]
    if "letter-block" in css_classes:
        paragraph.paragraph_format.space_after = Pt(LAYOUT["blockSpaceAfterPt"])
    elif "account-heading" in css_classes:
        paragraph.paragraph_format.space_after = Pt(LAYOUT["accountHeadingSpaceAfterPt"])
    elif "list-item" in css_classes:
        paragraph.paragraph_format.space_after = Pt(max(6, round(LAYOUT["paragraphSpaceAfterPt"] * 0.7)))
    else:
        paragraph.paragraph_format.space_after = Pt(LAYOUT["paragraphSpaceAfterPt"])

    if "list-item" in css_classes:
        paragraph.paragraph_format.left_indent = Inches(0.34)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)


def resolve_exhibit_image(block: Block, exhibits_dir: Optional[Path], warnings: List[str]) -> Optional[Path]:
    """Map an image block to a file inside the exhibits dir (basename only —
    figure markup carries data-exhibit-file). Missing assets warn-and-skip
    (memorandum semantics): the letter must never hard-fail on one lost PNG."""
    name = Path(str(block.image_src)).name
    if not name:
        return None
    if exhibits_dir is None:
        warnings.append(f"image block '{name}' skipped — no --exhibits-dir provided")
        return None
    if not PIL_AVAILABLE:
        warnings.append(f"image block '{name}' skipped — PIL unavailable")
        return None
    path = exhibits_dir / name
    if not path.exists():
        warnings.append(f"exhibit image missing on disk — slide skipped: {name}")
        return None
    return path


def build_docx(blocks: List[Block], output_path: Path, exhibits_dir: Optional[Path] = None, warnings: Optional[List[str]] = None):
    warnings = warnings if warnings is not None else []
    document = Document()
    configure_document(document)

    for block in blocks:
        if block.kind == "image":
            image_path = resolve_exhibit_image(block, exhibits_dir, warnings)
            if image_path is None:
                continue
            size = scaled_image_size(image_path)
            document.add_picture(str(image_path), width=Inches(size["width"]))
            picture_paragraph = document.paragraphs[-1]
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            picture_paragraph.paragraph_format.keep_with_next = True
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_runs_to_paragraph(paragraph, block.runs)
        apply_block_spacing(paragraph, block)

    document.save(str(output_path))


def runs_to_reportlab_markup(runs: List[InlineRun]) -> str:
    parts: List[str] = []
    for run in runs:
        if run.line_break:
            parts.append("<br/>")
            continue
        text = (
            run.text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        if run.underline:
            text = f"<u>{text}</u>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.bold:
            text = f"<b>{text}</b>"
        parts.append(text)
    return "".join(parts)


def build_pdf(blocks: List[Block], output_path: Path, exhibits_dir: Optional[Path] = None, warnings: Optional[List[str]] = None):
    if not REPORTLAB_AVAILABLE:
        return None
    warnings = warnings if warnings is not None else []

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=LAYOUT["marginInches"] * 72,
        rightMargin=LAYOUT["marginInches"] * 72,
        topMargin=LAYOUT["marginInches"] * 72,
        bottomMargin=LAYOUT["marginInches"] * 72,
    )
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "LetterBase",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=LAYOUT["fontSizePt"],
        leading=LAYOUT["fontSizePt"] * LAYOUT["lineHeight"],
        spaceAfter=LAYOUT["paragraphSpaceAfterPt"],
        textColor=colors.HexColor("#111827"),
        alignment=TA_LEFT,
    )

    caption_style = ParagraphStyle(
        "ExhibitCaption",
        parent=base,
        spaceBefore=4,
        spaceAfter=LAYOUT["blockSpaceAfterPt"],
    )

    story = []
    consumed = set()
    for idx, block in enumerate(blocks):
      if idx in consumed:
          continue
      if block.kind == "image":
          # Must precede the empty-markup skip below — image blocks carry no runs.
          image_path = resolve_exhibit_image(block, exhibits_dir, warnings)
          if image_path is None:
              continue
          size = scaled_image_size(image_path)
          flowable = RLImage(str(image_path), width=size["width"] * inch, height=size["height"] * inch)
          flowable.hAlign = "LEFT"
          bundle = [flowable]
          next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
          if next_block is not None and next_block.kind != "image" and "exhibit-caption" in set(next_block.css_class.split()):
              caption_markup = runs_to_reportlab_markup(next_block.runs)
              if caption_markup.strip():
                  bundle.append(Paragraph(caption_markup, caption_style))
              consumed.add(idx + 1)
          story.append(KeepTogether(bundle))
          continue
      markup = runs_to_reportlab_markup(block.runs)
      if not markup.strip():
          continue
      style = base
      css_classes = set(block.css_class.split())
      if "letter-block" in css_classes:
          style = ParagraphStyle(
              "LetterBlock",
              parent=base,
              spaceAfter=LAYOUT["blockSpaceAfterPt"],
          )
      elif "account-heading" in css_classes:
          style = ParagraphStyle(
              "LetterAccountHeading",
              parent=base,
              spaceAfter=LAYOUT["accountHeadingSpaceAfterPt"],
          )
      elif "list-item" in css_classes:
          style = ParagraphStyle(
              "LetterListItem",
              parent=base,
              leftIndent=0.34 * 72,
              firstLineIndent=-0.18 * 72,
              spaceAfter=max(6, round(LAYOUT["paragraphSpaceAfterPt"] * 0.7)),
          )
      story.append(Paragraph(markup, style))
    doc.build(story)
    return output_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("draft_json")
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--exhibits-dir", help="directory holding exhibit PNGs referenced by figure blocks (letter Mode A)")
    args = parser.parse_args()

    draft_path = Path(args.draft_json)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    exhibits_dir = Path(args.exhibits_dir) if args.exhibits_dir else None

    draft = json.loads(draft_path.read_text())
    html = draft.get("fullDocumentHtml") or draft.get("renderState", {}).get("previewHtml", "")
    blocks = parse_blocks(html)
    image_blocks = sum(1 for block in blocks if block.kind == "image")
    warnings: List[str] = []

    docx_path = output_dir / "dispute-letter.docx"
    build_docx(blocks, docx_path, exhibits_dir=exhibits_dir, warnings=warnings)

    pdf_path = output_dir / "dispute-letter.pdf"
    pdf_written = build_pdf(blocks, pdf_path, exhibits_dir=exhibits_dir, warnings=warnings)

    deduped_warnings = list(dict.fromkeys(warnings))
    print(json.dumps({
        "docxPath": str(docx_path),
        "pdfPath": str(pdf_written) if pdf_written else None,
        "reportlabAvailable": REPORTLAB_AVAILABLE,
        "imageBlocks": image_blocks,
        "warnings": deduped_warnings,
    }))


if __name__ == "__main__":
    main()

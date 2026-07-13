"""Memorandum of disputed items and supporting evidence (Exhibit Builder Phase 2).

Mode B companion document: a clean dispute letter goes to the bureau, and this
memorandum carries the evidence — one section per exhibit, numbered to match the
letter's dispute order exactly, each with its dispute statement, the screenshot(s),
and per-screenshot source-page citations.

Inputs: the draft JSON (identity/metadata) and an exhibits directory produced by
dispute_evidence_generator.py --exhibits-dir (exhibits-manifest.json + PNGs).
Outputs: evidence-memorandum.docx (always) and evidence-memorandum.pdf when
reportlab is available — mirroring dispute_letter_generator.py's behavior.
"""

import argparse
import json
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from dispute_letter_generator import (
    CONTENT_WIDTH_INCHES,
    LAYOUT,
    MAX_IMAGE_HEIGHT_INCHES,
    REPORTLAB_AVAILABLE,
    configure_document,
    scaled_image_size,
)

if REPORTLAB_AVAILABLE:
    from reportlab.lib.pagesizes import letter as letter_pagesize
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image as RLImage,
        KeepTogether,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

def escape_paragraph_text(text: object) -> str:
    """reportlab Paragraph parses XML-ish markup; report-derived text must be
    escaped (same idiom as dispute_letter_generator.runs_to_reportlab_markup)."""
    return (
        str(text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def format_entity(entity_key: str) -> str:
    """Human-readable entity heading. Account keys are 'name::number'; other entity
    types carry internal tokens that must never print raw in a mailed document."""
    parts = [part for part in str(entity_key or "").split("::") if part]
    if not parts:
        return ""
    kind = parts[0].strip().lower()
    if kind == "consumer_information_indicator":
        descriptor = parts[1] if len(parts) > 1 else ""
        return f"Consumer Information Indicator — {descriptor}".rstrip(" —") if descriptor else "Consumer Information Indicator"
    if kind == "public_record":
        descriptor = " — ".join(part for part in parts[1:3] if part and not part.isdigit())
        return f"Public Record — {descriptor}".rstrip(" —") if descriptor else "Public Record"
    if kind == "report":
        return "Credit Report (report-level dispute)"
    name = parts[0].upper()
    tail = parts[-1] if len(parts) > 1 and parts[-1] != parts[0] else ""
    return f"{name} — {tail}" if tail else name


def memorandum_header_lines(draft: Dict[str, object], exhibits_manifest: Dict[str, object]) -> List[str]:
    identity = draft.get("identity") or {}
    metadata = draft.get("metadata") or {}
    lines = []
    if identity.get("fullLegalName"):
        lines.append(f"Consumer: {identity['fullLegalName']}")
    if metadata.get("bureau"):
        lines.append(f"Credit bureau: {metadata['bureau']}")
    if metadata.get("reportNumber"):
        lines.append(f"Report / confirmation number: {metadata['reportNumber']}")
    if metadata.get("reportDate"):
        lines.append(f"Report date: {metadata['reportDate']}")
    if metadata.get("letterDate"):
        lines.append(f"Companion dispute letter dated: {metadata['letterDate']}")
    lines.append(f"Exhibits: {exhibits_manifest.get('exhibitCount') or 0}")
    return lines


INTRO_TEXT = (
    "This memorandum accompanies the referenced dispute letter and compiles the "
    "supporting evidence for each disputed item. Exhibit numbering matches the "
    "dispute numbering used in the letter. Each screenshot is reproduced from the "
    "consumer's credit report exactly as furnished, with the disputed information "
    "highlighted; source page numbers reference the enclosed highlighted copy of "
    "the full report."
)


def build_docx(
    draft: Dict[str, object],
    exhibits_manifest: Dict[str, object],
    exhibits_dir: Path,
    output_path: Path,
) -> List[str]:
    warnings: List[str] = []
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MEMORANDUM OF DISPUTED ITEMS AND SUPPORTING EVIDENCE")
    run.bold = True
    run.font.size = Pt(LAYOUT["fontSizePt"] + 2)

    for line in memorandum_header_lines(draft, exhibits_manifest):
        document.add_paragraph(line)
    document.add_paragraph(INTRO_TEXT)

    for index, exhibit in enumerate(exhibits_manifest.get("exhibits") or []):
        if index > 0:
            document.add_page_break()
        heading = document.add_paragraph()
        heading_run = heading.add_run(f"EXHIBIT {exhibit.get('exhibit')}")
        heading_run.bold = True
        heading_run.font.size = Pt(LAYOUT["fontSizePt"] + 1)

        entity = format_entity(str(exhibit.get("entityKey") or ""))
        if entity:
            entity_paragraph = document.add_paragraph()
            entity_paragraph.add_run(entity).bold = True
        issue_paragraph = document.add_paragraph()
        issue_paragraph.add_run(str(exhibit.get("issueLabel") or "")).bold = True
        if exhibit.get("reasonSummary"):
            document.add_paragraph(str(exhibit["reasonSummary"]))
        if exhibit.get("requestedAction"):
            document.add_paragraph(f"Requested action: {exhibit['requestedAction']}")

        for slide in exhibit.get("slides") or []:
            image_path = exhibits_dir / str(slide.get("file") or "")
            if not image_path.exists():
                warnings.append(f"exhibit {exhibit.get('exhibit')}: missing image {slide.get('file')}")
                continue
            size = scaled_image_size(image_path)
            document.add_picture(str(image_path), width=Inches(size["width"]))
            picture_paragraph = document.paragraphs[-1]
            picture_paragraph.paragraph_format.keep_with_next = True
            caption = document.add_paragraph()
            caption_text = f"Source: Credit report, page {slide.get('pageNumber')}"
            if slide.get("label"):
                caption_text += f" — {slide['label']}"
            caption_run = caption.add_run(caption_text)
            caption_run.italic = True
            caption_run.font.size = Pt(max(8, LAYOUT["fontSizePt"] - 2))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return warnings


def build_pdf(
    draft: Dict[str, object],
    exhibits_manifest: Dict[str, object],
    exhibits_dir: Path,
    output_path: Path,
) -> Optional[Path]:
    if not REPORTLAB_AVAILABLE:
        return None
    styles = getSampleStyleSheet()
    body = ParagraphStyle("MemoBody", parent=styles["Normal"], fontName="Times-Roman",
                          fontSize=LAYOUT["fontSizePt"], leading=LAYOUT["fontSizePt"] + 3)
    bold = ParagraphStyle("MemoBold", parent=body, fontName="Times-Bold")
    title_style = ParagraphStyle("MemoTitle", parent=bold, fontSize=LAYOUT["fontSizePt"] + 2,
                                 alignment=1, spaceAfter=12)
    caption_style = ParagraphStyle("MemoCaption", parent=body, fontName="Times-Italic",
                                   fontSize=max(8, LAYOUT["fontSizePt"] - 2), spaceAfter=10)

    flow = [Paragraph("MEMORANDUM OF DISPUTED ITEMS AND SUPPORTING EVIDENCE", title_style)]
    for line in memorandum_header_lines(draft, exhibits_manifest):
        flow.append(Paragraph(escape_paragraph_text(line), body))
    flow.append(Spacer(1, 8))
    flow.append(Paragraph(INTRO_TEXT, body))

    content_width = CONTENT_WIDTH_INCHES * inch
    for index, exhibit in enumerate(exhibits_manifest.get("exhibits") or []):
        flow.append(PageBreak() if index > 0 else Spacer(1, 12))
        flow.append(Paragraph(f"EXHIBIT {exhibit.get('exhibit')}", bold))
        entity = format_entity(str(exhibit.get("entityKey") or ""))
        if entity:
            flow.append(Paragraph(escape_paragraph_text(entity), bold))
        flow.append(Paragraph(escape_paragraph_text(exhibit.get("issueLabel")), bold))
        if exhibit.get("reasonSummary"):
            flow.append(Paragraph(escape_paragraph_text(exhibit["reasonSummary"]), body))
        if exhibit.get("requestedAction"):
            flow.append(Paragraph(escape_paragraph_text(f"Requested action: {exhibit['requestedAction']}"), body))
        flow.append(Spacer(1, 6))
        for slide in exhibit.get("slides") or []:
            image_path = exhibits_dir / str(slide.get("file") or "")
            if not image_path.exists():
                continue
            size = scaled_image_size(image_path)
            caption_text = f"Source: Credit report, page {slide.get('pageNumber')}"
            if slide.get("label"):
                caption_text += f" — {slide['label']}"
            flow.append(
                KeepTogether([
                    RLImage(str(image_path), width=size["width"] * inch, height=size["height"] * inch),
                    Paragraph(escape_paragraph_text(caption_text), caption_style),
                ])
            )

    document = SimpleDocTemplate(
        str(output_path), pagesize=letter_pagesize,
        leftMargin=LAYOUT["marginInches"] * inch, rightMargin=LAYOUT["marginInches"] * inch,
        topMargin=LAYOUT["marginInches"] * inch, bottomMargin=LAYOUT["marginInches"] * inch,
        title="Memorandum of Disputed Items and Supporting Evidence",
    )

    def draw_footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Times-Italic", 8)
        canvas.drawCentredString(
            letter_pagesize[0] / 2.0, 0.55 * inch, f"Page {canvas.getPageNumber()}"
        )
        canvas.restoreState()

    document.build(flow, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return output_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("draft_json")
    parser.add_argument("--exhibits-dir", required=True)
    parser.add_argument("--output-dir", required=True)
    args = parser.parse_args()

    draft = json.loads(Path(args.draft_json).read_text())
    exhibits_dir = Path(args.exhibits_dir)
    manifest_path = exhibits_dir / "exhibits-manifest.json"
    if not manifest_path.exists():
        print(json.dumps({"error": f"exhibits manifest not found: {manifest_path}"}))
        raise SystemExit(1)
    exhibits_manifest = json.loads(manifest_path.read_text())
    carried_warnings: List[str] = list(exhibits_manifest.get("warnings") or [])
    if not exhibits_manifest.get("exhibits"):
        carried_warnings.append("memorandum built with ZERO exhibits — not send-ready")
    for entry in exhibits_manifest.get("exhibits") or []:
        if not entry.get("slides"):
            carried_warnings.append(
                f"exhibit {entry.get('exhibit')}: no renderable evidence images"
            )

    output_dir = Path(args.output_dir)
    docx_path = output_dir / "evidence-memorandum.docx"
    warnings = build_docx(draft, exhibits_manifest, exhibits_dir, docx_path)
    pdf_path = build_pdf(draft, exhibits_manifest, exhibits_dir, output_dir / "evidence-memorandum.pdf")

    print(
        json.dumps(
            {
                "docxPath": str(docx_path),
                "pdfPath": str(pdf_path) if pdf_path else None,
                "exhibitCount": exhibits_manifest.get("exhibitCount") or 0,
                "reportlabAvailable": REPORTLAB_AVAILABLE,
                **({"warnings": carried_warnings + warnings} if (carried_warnings or warnings) else {}),
            }
        )
    )


if __name__ == "__main__":
    main()

from __future__ import annotations

import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("~/credit clarify - gain equity/3 Bureau Extractor")
ENGINE_PATH = ROOT / "src/features/dispute-letters/reasonEngine.ts"
OUTPUT_DIR = ROOT / "output/doc"
OUTPUT_PATH = OUTPUT_DIR / "dispute-engine-consumer-reason-reference.docx"


CATEGORY_LABELS = {
    "account_identity": "Identity",
    "payment_history": "Payment History",
    "balance_amount": "Balance / Amounts",
    "charge_off_collection": "Collection / Charge-Off",
    "date_reporting_timeline": "Status / Dates",
    "personal_information": "Personal Information",
    "attorney_escalation": "Attorney Escalation",
    "report_review": "Report Review",
}

NON_ACCOUNT_RULES = [
    {
        "issueType": "multiple_social_security_numbers",
        "issueLabel": "Multiple Social Security numbers reported",
        "category": "attorney_escalation",
        "description": "Flags report files that expose more than one Social Security number variation and routes them to attorney escalation.",
    },
    {
        "issueType": "personal_information_name_mismatch",
        "issueLabel": "Name mismatch",
        "category": "personal_information",
        "description": "Checks whether the provided consumer name conflicts with the name reflected in the report.",
    },
    {
        "issueType": "personal_information_address_mismatch",
        "issueLabel": "Address mismatch",
        "category": "personal_information",
        "description": "Checks whether the provided current mailing address conflicts with the addresses reflected in the report.",
    },
    {
        "issueType": "report_review_request",
        "issueLabel": "Report review requested",
        "category": "report_review",
        "description": "Fallback reason emitted when no other selected dispute reasons remain and the report still needs full reinvestigation review.",
    },
]

CATEGORY_INTROS = {
    "account_identity": "These reasons focus on whether the report identifies the account and the reporting company clearly enough to verify that the tradeline is accurate and belongs on the file.",
    "payment_history": "These reasons focus on whether the month-by-month payment timeline makes sense, is complete, and lines up with the rest of the account details.",
    "balance_amount": "These reasons focus on whether the reported balances, limits, high balances, and amount fields are complete and internally consistent.",
    "charge_off_collection": "These reasons focus on charge-off and collection reporting, especially when the negative status is not fully supported by the account history.",
    "date_reporting_timeline": "These reasons focus on whether the dates on the account line up with the payment history, balance history, and reported status.",
    "personal_information": "These reasons focus on whether the report identifies the consumer correctly in the personal-information section.",
    "attorney_escalation": "These issues are serious enough that the workflow routes them to attorney escalation rather than treating them like an ordinary account dispute.",
    "report_review": "This is a fallback review reason used when the report still needs a full reinvestigation even though no more specific selected reason is being exported.",
}

CONSUMER_SUMMARIES = {
    "duplicate_conflicting_tradeline": "The same account may be appearing more than once with conflicting details, such as different balances or statuses.",
    "missing_account_number": "The report does not show enough of the account number to verify that the tradeline is being reported accurately.",
    "missing_furnisher_identification": "The report does not clearly identify the company reporting the account because key contact details are missing.",
    "missing_account_status": "The account is negative or collection-related, but the report does not clearly show the current account status.",
    "incomplete_original_creditor_identity": "A collection-style account does not clearly identify the original creditor or lender behind the debt.",
    "responsibility_requires_special_handling": "The account is reported with a special responsibility type, such as joint or authorized user, so it needs extra review before it is disputed like a standard individual account.",
    "derogatory_status_without_monthly_support": "The account is being reported as late, past due, charged off, or otherwise negative, but the report does not show enough monthly history to support that negative status.",
    "payment_history_24_month_past_due_conflict": "The 24-month detail shows a past-due amount in months where the payment history looks current or blank, which can make the timeline inconsistent.",
    "payment_history_24_month_activity_conflict": "The 24-month account detail shows activity in months where the payment history table is blank, which can make the reporting look incomplete.",
    "amount_past_due_history_conflict": "The report shows positive past-due amounts without matching late-payment codes in the monthly payment history.",
    "delinquency_progression_inconsistency": "The late-payment sequence drops backward to a less severe level without first returning to current, such as 30 -> 60 -> 30 without an OK month in between.",
    "thirty_day_late_without_full_30_day_interval": "The report shows a 30-day-late month even though the strongest available date evidence suggests there was less than a full 30-day interval.",
    "severe_delinquency_jump_without_predecessor_support": "The payment history jumps too quickly into a more serious delinquency level, such as going from current or 30 days late straight to a much worse status without the expected step in between.",
    "reaging_jump_after_current_reset": "The account appears to return to current, but then jumps back to a severe late status instead of restarting at 30 days late.",
    "blank_gap_before_derogatory_month": "The report skips one or more months right before a late or derogatory month, which can make the timeline incomplete.",
    "retroactive_derogatory_backfill_after_reporting_gap": "A long reporting gap is followed by severe negative history, making it look like derogatory information may have been added back into the timeline later without full support.",
    "charge_off_or_collection_without_monthly_build_up": "The report shows charge-off or collection-style monthly history without enough earlier late-payment buildup to support that level of negative reporting.",
    "payment_plan_or_forbearance_context_without_history": "The report suggests the account was on a payment plan, deferment, or forbearance, but does not show enough payment history to explain what happened during that period.",
    "payment_plan_or_forbearance_context_with_derogatory_conflict": "The report suggests the account was on a payment plan, deferment, or forbearance, but the monthly history still shows a contradictory worsening delinquency pattern.",
    "payment_activity_conflicts_with_delinquency_progression": "The report shows signs that payments were being made, but the delinquency still worsens in a way that does not make sense.",
    "balance_reduction_conflicts_with_worsening_delinquency": "The balance goes down as if payments were made, but the late-payment history becomes more severe at the same time.",
    "closed_account_final_month_reporting_incomplete": "A closed account should still show activity for the month it closed. The report may be missing that final-month history.",
    "closed_account_actual_payment_conflicts_with_closure_month_history": "The amount reported as the last actual payment on a closed account does not match the payment information shown for the closure month.",
    "payment_history_balance_history_conflict": "The balance history does not line up with the payment activity the report is showing.",
    "balance_history_monthly_gap_conflict": "The report shows payment-history activity in months where the balance-history table is missing values.",
    "credit_limit_not_supported_by_history": "The credit limit being reported is not supported by the account's historical limit information.",
    "high_balance_not_supported_by_history": "The high balance being reported is not supported by the account's historical balance information.",
    "date_of_first_delinquency_conflict": "The reported date of first delinquency does not line up with the earliest late or derogatory month shown in the history.",
    "status_updated_timeline_conflict": "The status-updated date does not fit the payment-history timeline shown on the account.",
    "balance_updated_timeline_conflict": "The balance-updated date does not fit the balance-history timeline shown on the account.",
    "on_record_until_conflict": "The estimated removal date does not line up logically with the reported date of first delinquency.",
    "closed_account_missing_closure_timing": "The account is marked closed, but the report does not clearly show when it was closed.",
    "student_loan_lender_identity_mismatch": "The account appears to be a student loan, but the report does not clearly identify the lender or funding institution.",
    "report_review_request": "This is a general request for a full reinvestigation when the report still needs review but no more specific selected reason is being exported.",
}


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def parse_account_rules(engine_text: str) -> list[dict[str, str]]:
    pattern = re.compile(
        r'issueType:\s*"([^"]+)",\s*'
        r'issueLabel:\s*"([^"]+)",\s*'
        r'category:\s*"([^"]+)",\s*'
        r'description:\s*"([^"]+)"',
        re.S,
    )
    rules: list[dict[str, str]] = []
    seen: set[str] = set()
    for issue_type, issue_label, category, description in pattern.findall(engine_text):
        if issue_type in seen:
            continue
        seen.add(issue_type)
        rules.append(
            {
                "issueType": issue_type,
                "issueLabel": issue_label,
                "category": category,
                "description": description.replace("\\n", " ").strip(),
            }
        )
    return rules


def add_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    if "Small Note" not in document.styles:
        style = document.styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(9)


def fallback_consumer_summary(description: str) -> str:
    text = description.strip()
    if text.startswith("Checks whether "):
        text = text[len("Checks whether ") :]
    replacements = {
        "tradeline": "account",
        "discloses": "shows",
        "usable": "clear",
        "month-by-month": "month-by-month",
        "current-payment": "current-payment",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text[:1].upper() + text[1:] if text else text
    if not text.endswith("."):
        text += "."
    return text


def consumer_summary(rule: dict[str, str]) -> str:
    return CONSUMER_SUMMARIES.get(rule["issueType"], fallback_consumer_summary(rule["description"]))


def build_document(account_rules: list[dict[str, str]], non_account_rules: list[dict[str, str]]) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    add_base_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Possible Dispute Reasons the Engine Can Look For")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    subtitle = document.add_paragraph(style="Small Note")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f"Generated from the current dispute-rule catalog in reasonEngine.ts on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    )

    intro = document.add_paragraph()
    intro.add_run("Scope: ").bold = True
    intro.add_run(
        "This document explains, in plain language, the kinds of dispute reasons the engine can look for. "
        "A reason will only show up when the report contains enough evidence to evaluate it. "
        "The engine may mark a reason as Available, Clear, Not available, or Not applicable depending on the facts in the report."
    )

    summary = document.add_paragraph()
    summary.add_run("Counts: ").bold = True
    summary.add_run(
        f"{len(account_rules)} automated account rules, {len(non_account_rules)} automated non-account rules."
    )

    grouped_account_rules: dict[str, list[dict[str, str]]] = defaultdict(list)
    for rule in account_rules:
        grouped_account_rules[rule["category"]].append(rule)

    document.add_paragraph()
    document.add_paragraph("Account Dispute Reasons", style="Heading 1")

    for category in [
        "account_identity",
        "payment_history",
        "balance_amount",
        "charge_off_collection",
        "date_reporting_timeline",
    ]:
        rules = sorted(grouped_account_rules.get(category, []), key=lambda entry: entry["issueLabel"].lower())
        if not rules:
            continue

        document.add_paragraph(CATEGORY_LABELS.get(category, category), style="Heading 2")
        category_note = document.add_paragraph()
        category_note.add_run(CATEGORY_INTROS.get(category, "")).italic = True
        for rule in rules:
            rule_heading = document.add_paragraph()
            rule_heading.add_run(rule["issueLabel"]).bold = True
            summary_paragraph = document.add_paragraph(style="List Bullet")
            summary_paragraph.add_run(consumer_summary(rule))
        document.add_paragraph()

    document.add_paragraph("Non-Account Dispute Reasons", style="Heading 1")
    grouped_non_account_rules: dict[str, list[dict[str, str]]] = defaultdict(list)
    for rule in non_account_rules:
        grouped_non_account_rules[rule["category"]].append(rule)

    for category in ["personal_information", "attorney_escalation", "report_review"]:
        rules = grouped_non_account_rules.get(category, [])
        if not rules:
            continue
        document.add_paragraph(CATEGORY_LABELS.get(category, category), style="Heading 2")
        category_note = document.add_paragraph()
        category_note.add_run(CATEGORY_INTROS.get(category, "")).italic = True
        for rule in rules:
            paragraph = document.add_paragraph()
            paragraph.add_run(rule["issueLabel"]).bold = True
            detail = document.add_paragraph(style="List Bullet")
            detail.add_run(consumer_summary(rule))
        document.add_paragraph()

    document.add_paragraph("Operator-Entered Reasons", style="Heading 1")
    operator_note = document.add_paragraph()
    operator_note.add_run("Manual account reasons").bold = True
    operator_note.add_run(
        " can be added in the dispute-review UI for session-local account disputes, but they are not automatically detected by the engine and are therefore not included in the automated rule counts above."
    )

    return document


def main() -> None:
    engine_text = ENGINE_PATH.read_text(encoding="utf-8")
    account_rules = parse_account_rules(engine_text)
    document = build_document(account_rules, NON_ACCOUNT_RULES)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
